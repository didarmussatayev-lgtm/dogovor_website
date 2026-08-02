"""
Create a minimal soglasie_template.docx for development/testing.
Run from the backend/ directory:

    pip install python-docx
    python create_template.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path(__file__).parent / "app" / "templates" / "soglasie_template.docx"


def main():
    doc = Document()

    # Title
    title = doc.add_heading("СОГЛАСИЕ НА ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("И ПРОВЕДЕНИЕ МЕДИЦИНСКИХ ПРОЦЕДУР", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Meta
    doc.add_paragraph("Номер соглашения: {{ agreement_id }}")
    doc.add_paragraph("Дата: {{ date }}")
    doc.add_paragraph()

    # Client info
    doc.add_heading("Данные клиента", level=3)
    for label, placeholder in [
        ("ФИО", "{{ full_name }}"),
        ("Телефон", "{{ phone }}"),
        ("ИИН", "{{ iin }}"),
    ]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(placeholder)

    doc.add_paragraph()
    allergy_heading = doc.add_paragraph()
    allergy_heading.add_run("Аллергия / непереносимость:").bold = True
    doc.add_paragraph("{{ allergy }}")

    doc.add_paragraph()

    # Consent text
    doc.add_heading("Согласие", level=3)
    doc.add_paragraph(
        "Я, {{ full_name }}, настоящим даю своё добровольное согласие на "
        "обработку моих персональных данных в соответствии с Законом "
        "Республики Казахстан «О персональных данных и их защите» "
        "и на проведение медицинских процедур."
    )

    doc.add_paragraph()

    # Signature
    doc.add_heading("Подпись клиента", level=3)
    doc.add_paragraph("{{ signature }}")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Template saved to {OUTPUT}")


if __name__ == "__main__":
    main()
